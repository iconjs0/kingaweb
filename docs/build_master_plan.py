from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from datetime import date

OUT = "kingaweb/docs/KingaWeb_Product_and_Engineering_Master_Plan.docx"
GREEN = "12392F"
LIME = "B9F34A"
PALE = "EAF3E7"
INK = "10221C"
MUTED = "65716C"
LIGHT = "F3F5F1"
RED = "9B1C1C"
GOLD = "7A5A00"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

def font(run, size=11, bold=False, color=INK, italic=False, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size, run.bold, run.italic = Pt(size), bold, italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run

styles = doc.styles
normal = styles["Normal"]
normal.font.name, normal.font.size, normal.font.color.rgb = "Aptos", Pt(10.5), RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.16
for name, size, before, after, color in [
    ("Title", 31, 0, 7, INK), ("Subtitle", 14, 0, 12, MUTED),
    ("Heading 1", 18, 18, 9, GREEN), ("Heading 2", 14, 13, 6, GREEN),
    ("Heading 3", 11.5, 9, 4, MUTED)]:
    s = styles[name]
    s.font.name, s.font.size, s.font.color.rgb = "Aptos Display", Pt(size), RGBColor.from_string(color)
    s.font.bold = name != "Subtitle"
    s.paragraph_format.space_before, s.paragraph_format.space_after = Pt(before), Pt(after)
    s.paragraph_format.keep_with_next = True

for list_name in ("List Bullet", "List Number"):
    s = styles[list_name]
    s.font.name, s.font.size = "Aptos", Pt(10.5)
    s.paragraph_format.left_indent = Inches(.5)
    s.paragraph_format.first_line_indent = Inches(-.25)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.line_spacing = 1.16

if "Callout" not in styles:
    callout_style = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
else:
    callout_style = styles["Callout"]
callout_style.font.name, callout_style.font.size = "Aptos", Pt(10.5)
callout_style.paragraph_format.left_indent = Inches(.18)
callout_style.paragraph_format.right_indent = Inches(.18)
callout_style.paragraph_format.space_before = Pt(8)
callout_style.paragraph_format.space_after = Pt(10)
callout_style.paragraph_format.line_spacing = 1.15

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None: tcPr.append(shd)

def margins(cell, top=80, start=120, bottom=80, end=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if tcMar.getparent() is None: tcPr.append(tcMar)
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{key}")) or OxmlElement(f"w:{key}")
        node.set(qn("w:w"), str(val)); node.set(qn("w:type"), "dxa")
        if node.getparent() is None: tcMar.append(node)

def table(rows, widths, font_size=8.8, cell_y=80):
    t = doc.add_table(rows=1, cols=len(rows[0]))
    t.alignment, t.autofit = WD_TABLE_ALIGNMENT.LEFT, False
    tPr = t._tbl.tblPr
    tblW = tPr.find(qn("w:tblW")) or OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "9360"); tblW.set(qn("w:type"), "dxa")
    if tblW.getparent() is None: tPr.append(tblW)
    tblInd = OxmlElement("w:tblInd"); tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa"); tPr.append(tblInd)
    grid = t._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    dxa = [round(w * 1440) for w in widths]
    for w in dxa:
        node = OxmlElement("w:gridCol"); node.set(qn("w:w"), str(w)); grid.append(node)
    for ri, rowdata in enumerate(rows):
        row = t.rows[0] if ri == 0 else t.add_row()
        if ri == 0:
            trPr = row._tr.get_or_add_trPr(); rep = OxmlElement("w:tblHeader"); rep.set(qn("w:val"), "true"); trPr.append(rep)
        for ci, text in enumerate(rowdata):
            cell = row.cells[ci]; cell.width = Inches(widths[ci]); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW")) or OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(dxa[ci])); tcW.set(qn("w:type"), "dxa")
            if tcW.getparent() is None: cell._tc.get_or_add_tcPr().append(tcW)
            margins(cell, cell_y, 120, cell_y, 120); shade(cell, GREEN if ri == 0 else (LIGHT if ri % 2 == 0 else "FFFFFF"))
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.05
            font(p.add_run(str(text)), font_size, ri == 0, "FFFFFF" if ri == 0 else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def p(text="", bold_lead=None):
    par = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        font(par.add_run(bold_lead), bold=True)
        font(par.add_run(text[len(bold_lead):]))
    else: font(par.add_run(text))
    return par

def bullet(text):
    par = doc.add_paragraph(style="List Bullet"); font(par.add_run(text)); return par

def number(text):
    par = doc.add_paragraph(style="List Number"); font(par.add_run(text)); return par

def callout(label, text, fill=PALE):
    par = doc.add_paragraph(style="Callout")
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); pPr.append(shd)
    spacing = pPr.find(qn("w:spacing")) or OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "170"); spacing.set(qn("w:after"), "170")
    if spacing.getparent() is None: pPr.append(spacing)
    font(par.add_run(label.upper()+"  "), 9.5, True, GREEN if fill != GREEN else "FFFFFF")
    font(par.add_run(text), 10.5, False, INK if fill != GREEN else "FFFFFF")

def heading(text, level=1):
    doc.add_heading(text, level=level)

def page_break(): doc.add_page_break()

# Running furniture
header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
font(header.add_run("KINGAWEB  /  PRODUCT & ENGINEERING MASTER PLAN"), 8.5, True, MUTED)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(footer.add_run("Living document  •  v1.0  •  August 2026"), 8, False, MUTED)

# Cover
for _ in range(5): doc.add_paragraph()
tag = doc.add_paragraph(); tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(tag.add_run("DEFENSIVE SECURITY  •  BUILT FOR AFRICA"), 10, True, GREEN)
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(title.add_run("KingaWeb"), 36, True, INK, name="Aptos Display")
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(sub.add_run("Product & Engineering Master Plan"), 18, False, GREEN, name="Aptos Display")
strap = doc.add_paragraph(); strap.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(strap.add_run("A living blueprint from prototype to trusted African security platform"), 12, False, MUTED, True)
doc.add_paragraph()
callout("North-star promise", "KingaWeb helps organizations discover public-facing risks early, understand them in plain language, and close them through a verifiable remediation workflow.", GREEN)
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(meta.add_run("Document owner: KingaWeb Product Team\nVersion 1.0  |  23 August 2026  |  Status: Approved working baseline"), 9.5, False, MUTED)
page_break()

heading("How to use this master plan", 1)
p("This document is the durable decision record for KingaWeb. It defines what we are building, why it should exist, how it must be engineered, and what “complete” means for each release. It should evolve through controlled versions as customer evidence, regulation, threat patterns and technology change.")
callout("Important distinction", "KingaWeb can reach 100% of a defined release, but a security platform is never permanently finished. After Version 1 reaches its release criteria, development continues through monitoring improvements, new integrations, threat research and regional expansion.")
heading("Decision hierarchy", 2)
for x in [
    "Safety and authorization outrank feature speed.",
    "Verified user problems outrank attractive but unsupported ideas.",
    "Reliable, explainable findings outrank a large count of noisy findings.",
    "Mobile usability and low-bandwidth operation are first-class requirements.",
    "A release is complete only when its acceptance, security, operational and documentation gates pass."]:
    bullet(x)
heading("Document governance", 2)
table([
    ["Item", "Rule"],
    ["Change control", "Major scope or architecture changes require a dated decision record."],
    ["Review cadence", "Review monthly during build; quarterly after public launch."],
    ["Versioning", "Major.Minor format; major for strategy/architecture changes, minor for clarified execution."],
    ["Source of truth", "This document plus repository issues, ADRs and release checklists."],
    ["Definition of done", "Code, tests, security checks, observability, documentation and acceptance evidence."]], [1.55,4.95])

page_break()
heading("Contents", 1)
for i, item in enumerate([
    "Executive vision and product thesis", "Market position and differentiation", "Users and jobs to be done",
    "Product scope and capability map", "Experience and interface system", "System architecture",
    "Scanner, findings and AI intelligence", "Security, privacy and abuse prevention", "Data model and APIs",
    "Delivery roadmap", "Engineering standards and testing", "Infrastructure and operations",
    "Commercial model and go-to-market", "Metrics, risks and governance", "Release-one completion checklist",
    "Post-launch continuous development", "Reference standards and glossary"], 1):
    p(f"{i:02d}  {item}")
page_break()

heading("1. Executive vision and product thesis", 1)
heading("1.1 Vision", 2)
p("KingaWeb will become the trusted, accessible external security companion for African small and growing organizations: a platform that continuously maps what the internet can see, translates technical exposure into business impact, and guides owners or their service partners from detection to verified resolution.")
heading("1.2 The problem", 2)
p("Many organizations depend on websites, cloud services, email domains and third-party platforms without maintaining an accurate view of their public attack surface. Existing enterprise tools are often expensive, technically dense or designed for mature security teams. A simple scanner alone is not enough: it produces a moment-in-time list, while risk changes whenever certificates, DNS, code, hosting or suppliers change.")
heading("1.3 Product thesis", 2)
for x in [
    "Continuous visibility is more valuable than a one-time grade.",
    "Owners act faster when a finding explains business impact, ownership, priority and a tested fix.",
    "Local delivery channels—especially mobile and WhatsApp-ready communication—can turn security data into action.",
    "Agencies and managed-service partners are a scalable distribution channel because they already maintain client websites.",
    "Regional relevance must come from workflow, language, affordability and support—not from unsupported claims that African threats are entirely different."]:
    bullet(x)
heading("1.4 Mission and boundaries", 2)
p("Mission: make responsible, continuous external security monitoring understandable and actionable for organizations that do not have a security operations centre.")
p("Boundary: KingaWeb is a defensive monitoring and remediation platform. It is not an anonymous attack tool, an unrestricted penetration-testing service, a guarantee that a system is secure, or a replacement for a professional assessment where risk is high.")
heading("1.5 North-star outcome", 2)
callout("North-star metric", "Verified risk-reduction events per protected organization per month: a previously confirmed issue is fixed, automatically rechecked and closed with evidence.")

heading("2. Market position and differentiation", 1)
p("Public research shows relevant African cybersecurity consultancies and products already exist, including offerings for vulnerability assessment, website monitoring and SME subscriptions. KingaWeb must therefore earn distinction through execution and product depth; “first in Africa” will not be used without independently verifiable evidence.")
heading("2.1 Positioning statement", 2)
callout("Position", "For African SMEs and the agencies that support them, KingaWeb is a mobile-first external security management platform that turns continuous technical signals into prioritized, bilingual remediation workflows and proof of improvement.")
heading("2.2 Differentiation pillars", 2)
table([
    ["Pillar", "What KingaWeb does differently", "Evidence of success"],
    ["Action, not alarm", "Each finding has owner, fix steps, due date, verification and evidence.", "Higher verified closure rate; lower time-to-remediate."],
    ["African operating fit", "Low-bandwidth UI, TZS/local pricing, Swahili/English, mobile-first alerts.", "Strong mobile completion and regional activation."],
    ["Agency operating system", "Multi-client workspaces, branded reports, delegated access and portfolio risk view.", "Agencies adopt and retain multiple clients."],
    ["Trust through clarity", "Explainable scoring, evidence timestamps, confidence levels and scan limitations.", "Low dispute/false-positive rate."],
    ["Responsible by design", "Ownership verification, safe scan profiles, audit logs and abuse controls.", "No unauthorized-scan incidents; passed reviews."],
    ["From web to perimeter", "Domains, DNS, TLS, email posture, exposed assets, impersonation and change monitoring.", "Growing assets protected per customer."]], [1.15,3.25,2.10])
heading("2.3 Initial wedge and expansion", 2)
p("The entry product is website, TLS, DNS and email-domain posture monitoring. The expansion is a governed external attack-surface platform covering subdomains, approved services, public cloud exposures, third-party assets, impersonation signals and agency-managed remediation.")

heading("3. Users and jobs to be done", 1)
table([
    ["Persona", "Primary job", "Key anxiety", "Success"],
    ["Business owner", "Know whether the digital business is exposed and what to do next.", "Technical reports are confusing or unaffordable.", "Clear priorities and proof that fixes worked."],
    ["Web agency", "Protect many client sites without repetitive manual checking.", "A client incident damages trust.", "Portfolio view, reports and fast remediation."],
    ["Developer/IT lead", "Detect configuration drift before customers or attackers do.", "No time to watch every external signal.", "Useful alerts, APIs and evidence."],
    ["Security consultant", "Deliver repeatable assessments and track remediation.", "Tool noise and fragmented evidence.", "Controlled scanning and client-ready reporting."],
    ["Organization executive", "Understand business exposure and trend without security jargon.", "Unknown risk and unclear accountability.", "Risk trend, ownership and audit trail."]], [1.15,2.25,1.55,1.55])
heading("3.1 Core journeys", 2)
for x in [
    "Owner: create workspace → add domain → verify ownership → first scan → understand top three actions → assign or share → verify fix.",
    "Agency: create portfolio → invite team → add clients → verify assets → review cross-client risk → generate branded report → track remediation.",
    "Developer: receive alert → inspect evidence and request trace → apply configuration change → trigger safe rescan → close finding.",
    "Executive: open monthly digest → view risk trend and critical open items → confirm responsible owner and due date."]:
    number(x)

heading("4. Product scope and capability map", 1)
heading("4.1 Capability domains", 2)
table([
    ["Domain", "Release-one capability", "Later evolution"],
    ["Asset inventory", "Verified domains, websites and DNS snapshots.", "Subdomains, IP ranges, SaaS and cloud assets."],
    ["Web posture", "HTTPS, TLS, headers, redirects, cookies and exposed metadata.", "Safe crawling, API posture and supply-chain signals."],
    ["Email posture", "SPF, DKIM discovery, DMARC and MX health.", "BIMI, spoof simulations and domain abuse monitoring."],
    ["Change detection", "Certificate, DNS, content hash and header changes.", "Risk-aware diffs and maintenance windows."],
    ["Remediation", "Guidance, ownership, status, evidence and rescan.", "Ticketing integrations and verified automation."],
    ["Communication", "Email and WhatsApp-ready share links/digests.", "Provider integrations, SMS and webhooks."],
    ["Reporting", "Current posture and trend PDF/CSV exports.", "Compliance mappings and executive portfolios."],
    ["Platform", "Workspaces, roles, audit log, billing-ready plans.", "Partner marketplace and public API ecosystem."]], [1.2,2.55,2.75])
heading("4.2 Explicit non-goals for Release One", 2)
for x in [
    "Exploit execution, password attacks, denial-of-service testing or aggressive unauthenticated port scanning.",
    "Endpoint antivirus, full SIEM, employee surveillance or packet interception.",
    "Guaranteeing security or issuing legal/compliance certification.",
    "AI-generated findings without deterministic evidence and a documented rule.",
    "Native mobile applications before the responsive PWA experience proves insufficient."]:
    bullet(x)

heading("5. Experience and interface system", 1)
heading("5.1 Experience principles", 2)
for x in [
    "Calm, credible and precise: security urgency without fear-based marketing.",
    "Progressive disclosure: business meaning first, technical evidence second.",
    "Mobile-first: every critical journey works at 360 px width and on constrained networks.",
    "Accessible: WCAG 2.2 AA target, keyboard navigation, visible focus and non-color status cues.",
    "Bilingual architecture: English and Kiswahili strings externalized from the first production build.",
    "No dark patterns: transparent trials, pricing, data retention and cancellation."]:
    bullet(x)
heading("5.2 Information architecture", 2)
table([
    ["Area", "Primary screens"],
    ["Public", "Landing, free check, methodology, trust centre, pricing, legal, status."],
    ["Onboarding", "Account, workspace, domain verification, first scan, result walkthrough."],
    ["Operate", "Overview, assets, findings, scans, remediation, reports, alerts."],
    ["Agency", "Client portfolio, team, assignments, white-label report configuration."],
    ["Administration", "Roles, integrations, billing, audit log, data export/deletion."],
    ["Internal", "Rules, scan queues, abuse review, support, service health, feature flags."]], [1.45,5.05])
heading("5.3 Dashboard hierarchy", 2)
p("The dashboard leads with: critical change since last scan; top three actions; overall trend; protected assets; recent scans; then deeper categories. A security score is never shown without its components, timestamp, confidence and limitations.")
heading("5.4 Design system", 2)
p("Visual identity uses deep forest green (#12392F), lime signal (#B9F34A), warm neutral surfaces and high-contrast ink. Components include severity badges, evidence cards, asset chips, trend charts, remediation timelines, verified-status markers, skeleton states and empty-state education. Storybook or an equivalent component catalogue becomes mandatory before public beta.")

heading("6. System architecture", 1)
callout("Architecture principle", "Separate the customer-facing control plane from the scanning data plane. Scanners are untrusted workers with minimal credentials, strict egress policy, bounded execution and no direct access to application secrets.")
heading("6.1 Target architecture", 2)
p("Browser/PWA → CDN and web application → API gateway → application services → PostgreSQL. The application publishes signed scan jobs to a queue. Isolated workers execute approved checks through controlled egress, store normalized evidence in object storage and results in PostgreSQL, then emit finding and notification events. A scheduler creates recurring jobs; an observability stack receives logs, metrics and traces.")
table([
    ["Layer", "Recommended direction", "Reason"],
    ["Web", "Next.js with TypeScript", "Strong server rendering, routing, PWA support and shared types."],
    ["API", "NestJS/TypeScript or FastAPI/Python", "Structured modules, validation and OpenAPI; choose one at foundation milestone."],
    ["Scanner", "Python workers", "Mature networking, DNS, TLS and security-analysis libraries."],
    ["Data", "Managed PostgreSQL + Redis-compatible queue/cache", "Relational integrity, JSON evidence and reliable jobs."],
    ["Artifacts", "S3-compatible encrypted object storage", "Reports and bounded evidence retention."],
    ["Identity", "Managed OIDC initially", "MFA, recovery and reduced authentication risk."],
    ["Infrastructure", "Containers + infrastructure as code", "Repeatability and future cloud portability."],
    ["Observability", "OpenTelemetry + managed logs/metrics/errors", "Vendor-neutral instrumentation and operational visibility."]], [1.2,2.3,3.0])
heading("6.2 Architectural qualities", 2)
for x in [
    "Tenant isolation at every query and object-storage path.",
    "Idempotent scan jobs with immutable job IDs and retry limits.",
    "Versioned rules, scores, APIs and evidence schemas.",
    "Horizontal worker scaling independent of web traffic.",
    "Graceful degradation: dashboards remain usable during scan backlog or notification outage.",
    "Regional deployment readiness without premature multi-region complexity."]:
    bullet(x)
heading("6.3 Architecture decisions required before implementation", 2)
for x in ["Cloud/hosting provider and data region", "Primary backend framework", "Queue implementation", "Identity provider", "Notification providers", "Payment provider", "Retention policy and Tanzania legal review"]:
    bullet(x)

heading("7. Scanner and findings engine", 1)
heading("7.1 Safe Release-One checks", 2)
table([
    ["Category", "Checks", "Default intensity"],
    ["Reachability", "Status, redirect chain, latency and uptime.", "One bounded request plus validation."],
    ["TLS", "Trust chain, protocol, hostname, expiry and selected configuration.", "One handshake; no exhaustive cipher flood."],
    ["HTTP", "Security headers, cookies, server disclosure, mixed-content indicators.", "HEAD/GET with response-size ceiling."],
    ["DNS", "A/AAAA, CNAME, NS, MX, CAA and changes.", "Cached, rate-limited lookups."],
    ["Email", "SPF, DMARC and DKIM-selector workflow.", "Passive DNS checks."],
    ["Content", "Unexpected change hash and approved defacement indicators.", "Small bounded sample; no personal content retention."],
    ["Domain", "Expiry where registry data permits; nameserver change.", "Passive registry/RDAP data."],
    ["Technology", "Conservative public fingerprinting with confidence.", "No intrusive probing."]], [1.15,3.65,1.70])
heading("7.2 Finding contract", 2)
p("Every rule produces: stable rule ID; rule version; asset; observed time; severity; confidence; evidence; business impact; remediation; verification method; standards mappings; and false-positive feedback. Findings are deduplicated across scans and transition through Open, Acknowledged, In Progress, Resolved, Accepted Risk, False Positive and Reopened.")
heading("7.3 Scoring", 2)
p("The score is a transparent weighted posture indicator, not a security guarantee. Critical findings cap the score; confidence affects presentation, not hidden multiplication. Rule weights and scoring versions are public. Historical scores remain tied to the scoring version used at the time.")
heading("7.4 False-positive discipline", 2)
for x in [
    "A rule needs fixtures for positive, negative and ambiguous states before release.",
    "Low-confidence signals are labeled informational until corroborated.",
    "Customers can dispute findings; dispositions feed rule review, not automatic suppression across tenants.",
    "Rule quality dashboard tracks confirmation rate, disputes and regression failures."]:
    bullet(x)

heading("7.5 AI-powered real-time intelligence", 2)
p("KingaWeb’s intelligence layer turns trusted observations into timely decisions. “Real-time” means event-driven processing immediately after an authorized change, scan result, provider signal or customer action; it does not mean intercepting private network traffic or continuously attacking targets. Deterministic rules remain the source of truth for findings. AI correlates, explains, prioritizes and assists.")
heading("7.5.1 Initial AI capabilities", 3)
table([
    ["Capability", "Customer value", "Required guardrail"],
    ["Risk Copilot", "Answers questions about the customer’s own assets, evidence and fixes in English or Kiswahili.", "Retrieval only from authorized workspace data; citations to evidence."],
    ["Change intelligence", "Summarizes what changed across TLS, DNS, headers, technology and content signals.", "No claim beyond observed diffs; confidence displayed."],
    ["Alert prioritization", "Groups related events, suppresses duplicates and recommends urgency using business context.", "Deterministic severity remains visible; user can override."],
    ["Remediation assistant", "Produces stack-specific fix instructions, validation steps and draft tickets.", "Human approval; never executes production changes in Release One."],
    ["Incident timeline", "Builds a readable timeline from scans, changes, alerts and remediation actions.", "Every statement links to timestamped source evidence."],
    ["Anomaly detection", "Flags unusual deviations from an asset’s historical baseline.", "Starts informational; promotion requires evaluation and corroboration."],
    ["Executive digest", "Converts posture changes into business-language daily/weekly summaries.", "No invented causes, guarantees or compliance claims."],
    ["Feedback learning", "Uses confirmed fixes and false-positive feedback to improve ranking.", "Tenant isolation; controlled evaluation before model/rule changes."]], [1.35,3.25,1.90], font_size=8.1, cell_y=55)
heading("7.5.2 Intelligence event pipeline", 3)
p("Authorized signal → immutable observation → deterministic rule evaluation → event stream → correlation window → anomaly/risk models → evidence-grounded explanation → policy gate → notification or recommended action → human decision → outcome feedback. Every stage carries tenant, asset, rule/model version, timestamp, confidence and correlation identifiers.")
heading("7.5.3 AI architecture", 3)
for x in [
    "Feature service computes bounded security features from normalized observations; raw untrusted page content is excluded by default.",
    "Correlation service groups related DNS, certificate, header, uptime and asset events within defined windows.",
    "Model gateway provides provider abstraction, redaction, model allowlists, budgets, timeouts and full request metadata without logging sensitive prompts by default.",
    "Retrieval layer searches only workspace-authorized evidence and approved remediation knowledge using tenant-scoped indexes.",
    "Policy engine validates outputs, blocks unsafe actions and requires human approval for notifications or integrations with meaningful impact.",
    "Evaluation service runs golden datasets, adversarial prompts, hallucination checks, bilingual quality tests and cost/latency benchmarks before rollout.",
    "AI telemetry records model/version, retrieved evidence IDs, confidence, policy decisions, user feedback and outcome—without exposing secrets."]:
    bullet(x)
heading("7.5.4 AI safety requirements", 3)
for x in [
    "Treat websites, DNS text, tickets and user content as untrusted data, never as instructions to the model.",
    "Prevent prompt injection through strict context separation, input labeling, output validation and tool allowlists.",
    "No unrestricted agent loops, shell access, scanner control or network access from the language model.",
    "Authorization is enforced by application services; an AI response can never grant access.",
    "High-impact actions require explicit human confirmation and are executed by deterministic downstream services.",
    "Customers can disable generative features and request deletion under the documented retention policy.",
    "Model providers must not train on customer data unless a separate explicit agreement permits it.",
    "AI features launch behind flags with rollback, cost limits, safety monitoring and incident procedures."]:
    bullet(x)
heading("7.5.5 AI delivery sequence", 3)
p("Private MVP: rule-based prioritization and evidence-grounded scan summaries. Private beta: Risk Copilot, change intelligence and draft remediation. Public V1: evaluated bilingual summaries, executive digests and feedback capture. V1.x: historical anomaly detection and cross-signal correlation. V2: carefully governed agent-assisted workflows, integrations and impersonation intelligence—still human-approved for consequential action.")

heading("8. Security, privacy and abuse prevention", 1)
heading("8.1 Threat model", 2)
p("Primary threats include unauthorized target scanning, server-side request forgery, cross-tenant access, compromised worker containers, queue poisoning, credential theft, report leakage, malicious domains, denial of wallet through expensive scans, dependency compromise, insider misuse and manipulated evidence.")
heading("8.2 Required controls", 2)
table([
    ["Control area", "Mandatory baseline"],
    ["Authorization", "Verify domain by DNS TXT, hosted file or meta tag before recurring or expanded scans."],
    ["SSRF", "Resolve before request; reject private, loopback, link-local, reserved and metadata addresses; revalidate every redirect and connection."],
    ["Isolation", "Ephemeral non-root workers, read-only filesystem, resource limits and no application-secret access."],
    ["Egress", "Network policy, approved protocols/ports, destination validation, timeouts and response limits."],
    ["Identity", "MFA for privileged roles, secure recovery, session rotation and step-up authentication."],
    ["Tenancy", "Workspace-scoped authorization, database row policies/guards, object-path isolation and tests."],
    ["Secrets", "Managed secret store, rotation, least privilege and no secrets in code/logs."],
    ["Audit", "Append-only security events for roles, assets, scans, exports, billing and administrative action."],
    ["Supply chain", "Pinned dependencies, SBOM, secret scanning, SAST, dependency alerts and signed releases."],
    ["Data", "Encryption in transit/at rest, minimization, retention schedule, export and deletion workflow."],
    ["Abuse", "Rate limits, quotas, reputation signals, manual review and immediate kill switch."],
    ["Disclosure", "Published vulnerability disclosure policy and response SLA before beta."]], [1.2,5.3])
heading("8.3 Secure development standard", 2)
p("Release One targets OWASP ASVS Level 2-aligned verification where applicable, OWASP Web Security Testing Guide coverage for relevant controls, and secure-by-design defaults. NIST CSF 2.0 informs KingaWeb’s customer-facing risk lifecycle: Govern, Identify, Protect, Detect, Respond and Recover.")
p("AI features additionally follow the NIST AI Risk Management Framework’s Govern, Map, Measure and Manage functions and current OWASP GenAI guidance. Evaluation covers prompt injection, sensitive-data disclosure, insecure output handling, excessive agency, model/provider dependency, denial-of-wallet and overreliance.")
heading("8.4 Privacy principles", 2)
for x in [
    "Collect public technical evidence only when needed; avoid page bodies and personal data by default.",
    "Separate operational telemetry from customer scan evidence.",
    "Provide retention choices, export and deletion; document exceptions for fraud/security records.",
    "Do not sell customer data or use private evidence to train models without explicit agreement.",
    "Complete a Tanzania data-protection/legal review before public processing of customer information."]:
    bullet(x)

heading("9. Data model and APIs", 1)
heading("9.1 Core entities", 2)
table([
    ["Entity", "Purpose", "Key relationships"],
    ["User", "Person and authentication identity.", "Memberships, assignments, audit events."],
    ["Workspace", "Tenant and billing boundary.", "Members, clients, assets, plans."],
    ["Client", "Agency-managed customer boundary.", "Workspace and assets."],
    ["Asset", "Verified domain/site/service.", "Verification, scans, findings."],
    ["Verification", "Ownership challenge and evidence.", "Asset, requester and expiry."],
    ["ScanJob", "Requested work and policy snapshot.", "Asset, profile and worker run."],
    ["Observation", "Immutable normalized evidence.", "Scan job and rule evaluation."],
    ["Finding", "Deduplicated risk lifecycle.", "Asset, rule, observations, owner."],
    ["Remediation", "Action, due date and verification trail.", "Finding and assignee."],
    ["Notification", "Delivery intent/result.", "User, event and channel."],
    ["AuditEvent", "Security/accountability record.", "Actor, workspace, target."],
    ["RuleVersion", "Detection and scoring logic metadata.", "Observations and findings."]], [1.15,2.75,2.60])
heading("9.2 API principles", 2)
for x in [
    "REST/OpenAPI first; resource-oriented endpoints under /v1.",
    "Opaque identifiers; cursor pagination; ISO 8601 UTC timestamps.",
    "Idempotency keys on job creation, billing and high-impact mutations.",
    "Consistent problem-details errors with request correlation IDs.",
    "Strict validation, content-size limits and explicit authorization middleware.",
    "Signed, replay-resistant webhooks with retry and delivery history."]:
    bullet(x)
heading("9.3 Initial endpoint groups", 2)
p("/auth and identity callbacks; /workspaces; /memberships; /clients; /assets; /verifications; /scans; /findings; /remediations; /reports; /notifications; /integrations; /audit-events; /billing; /webhooks. Administrative endpoints live on a separate internal surface.")

heading("10. Delivery roadmap", 1)
callout("Planning rule", "Dates are forecasts set after discovery and capacity review. Scope gates are fixed; calendar promises are not made before the team can estimate from evidence.")
table([
    ["Phase", "Outcome", "Exit gate"],
    ["0 — Discovery", "10–20 interviews; target segment and willingness-to-pay evidence.", "Problem brief, risk register and design partner."],
    ["1 — Foundation", "Production repo, design system, CI, architecture skeleton and threat model.", "ADR set, staging environment and security baseline."],
    ["2 — Private MVP", "Verified assets, safe scans, findings and manual rescan for 3–5 partners.", "Useful results; no critical unresolved product-security issue."],
    ["3 — Private beta", "Accounts, scheduling, history, alerts, remediation and agency workspace.", "10 active organizations; reliability and feedback gates."],
    ["4 — Public V1", "Billing, reports, trust centre, support and production operations.", "All Release-One completion criteria pass."],
    ["5 — V1.x", "Rule quality, Kiswahili, PWA, integrations and commercial optimization.", "Retention and remediation metrics improve."],
    ["6 — V2 perimeter", "Subdomain discovery, approved service exposure and impersonation workflows.", "Expanded authorization and legal/safety review."],
    ["7 — Regional scale", "Partner channel, regional payments, multilingual support and resilience.", "Repeatable acquisition and operational capacity."]], [1.2,3.4,1.9])
heading("10.1 Phase 0 — discovery deliverables", 2)
for x in [
    "Interview script and consent statement; at least five businesses, five agencies and three developers/IT leads.",
    "Competitor matrix based on verified capabilities, pricing, audience and workflow—not marketing claims.",
    "Clickable mobile prototype tested with at least five target users.",
    "One design partner agreeing to verify assets and review findings regularly.",
    "Decision record: exact first segment, pricing hypothesis and Release-One boundary."]:
    bullet(x)
heading("10.2 Foundation epics", 2)
for x in ["Monorepo and development environments", "Identity and tenant model", "Design tokens and component library", "Queue and isolated worker proof", "Rule SDK and fixtures", "Observability baseline", "Infrastructure as code", "Threat model and abuse controls"]:
    bullet(x)
heading("10.3 Release sequencing rule", 2)
p("A feature moves through Discovery → Designed → Threat-modeled → Implemented → Tested → Documented → Observed in staging → Released behind a flag → Measured. Work is not complete at code merge.")

heading("11. Engineering standards and testing", 1)
heading("11.1 Repository standards", 2)
for x in [
    "Protected main branch, reviewed pull requests and conventional commit/release notes.",
    "Format, lint, type-check, unit and security checks required in CI.",
    "Architecture Decision Records for durable choices; RFCs for cross-cutting changes.",
    "Feature flags for risky or incomplete behavior; no environment-specific code branches.",
    "Generated database migrations reviewed and rehearsed with rollback/forward strategy."]:
    bullet(x)
heading("11.2 Test pyramid", 2)
table([
    ["Level", "Coverage expectation"],
    ["Unit", "Rules, scoring, validation, permissions and state transitions."],
    ["Component", "API modules, repositories, queue behavior and notification rendering."],
    ["Integration", "Database isolation, object storage, DNS/TLS fixtures and worker sandbox."],
    ["Contract", "OpenAPI compatibility, webhook signatures and event schemas."],
    ["End-to-end", "Critical onboarding, verification, scan, remediation, billing and deletion journeys."],
    ["Security", "SAST, dependency/secret/container scanning, authorization tests and targeted DAST."],
    ["Resilience", "Timeouts, retry storms, worker crashes, queue backlog and provider outages."],
    ["Accessibility", "Automated checks plus keyboard, screen-reader and contrast review."],
    ["Performance", "Mobile page budgets, API latency and controlled scan throughput."]], [1.25,5.25])
heading("11.3 Quality targets for public V1", 2)
for x in [
    "No open Critical or High product-security vulnerabilities; Medium issues have owners and deadlines.",
    "99.9% monthly control-plane availability target after stabilization; status page reports incidents.",
    "P95 standard API latency under 500 ms excluding scan execution and third-party delays.",
    "Core web experience meets agreed mobile performance budget on representative constrained connection.",
    "All critical journeys pass keyboard and accessibility review.",
    "Backup restoration demonstrated, not merely configured."]:
    bullet(x)

heading("12. Infrastructure and operations", 1)
heading("12.1 Environments", 2)
p("Local development uses safe fixtures and a local mock target. Preview environments run per pull request with no production secrets. Staging mirrors production topology with synthetic data. Production access is least privilege, MFA protected, time bounded where possible and fully audited.")
heading("12.2 Deployment", 2)
for x in [
    "Build immutable images; generate SBOM; scan and sign release artifacts.",
    "Deploy database-compatible changes before application changes; use expand/contract migrations.",
    "Canary or gradual rollout for scanner rules and high-impact services.",
    "Automatic rollback only for well-understood health signals; otherwise operator-controlled rollback.",
    "Publish customer-visible incidents and material rule/scoring changes."]:
    bullet(x)
heading("12.3 Observability and SLOs", 2)
table([
    ["Service indicator", "Initial objective", "Alerting principle"],
    ["API availability", "99.9% monthly after public launch", "Page on sustained error-budget burn."],
    ["Scan completion", "95% of routine scans finish within 10 minutes", "Alert on backlog age and failure cluster."],
    ["Notification delivery", "99% accepted by provider within 5 minutes", "Alert on provider or template failures."],
    ["Evidence durability", "No acknowledged loss of committed scan evidence", "Alert on storage/write verification."],
    ["Queue health", "Oldest routine job below agreed threshold", "Scale workers or pause intake safely."],
    ["Rule quality", "Monitored dispute and confirmation rates", "Disable/regress noisy rule versions."]], [1.55,2.3,2.65])
heading("12.4 Incident response", 2)
p("Maintain severity definitions, on-call ownership, containment playbooks, customer communication templates, forensic logging, legal escalation contacts and post-incident review. Conduct one tabletop exercise before public launch and at least twice yearly thereafter.")

heading("13. Commercial model and go-to-market", 1)
heading("13.1 Plans hypothesis", 2)
table([
    ["Plan", "Audience", "Initial shape"],
    ["Free", "Owner evaluating posture", "One verified domain, limited manual checks, current snapshot."],
    ["Starter", "Small business", "Daily monitoring, history, alerts and remediation workflow."],
    ["Business", "Growing organization", "More assets/users, reports, integrations and priority support."],
    ["Agency", "Web/IT/security provider", "Client portfolio, delegated roles, branded reports and volume pricing."],
    ["Enterprise/Institution", "Complex organization", "Custom scope, SSO, retention, assurance and support agreement."]], [1.25,1.85,3.40])
p("Pricing remains a hypothesis until discovery measures willingness to pay and support cost. Security essentials such as MFA, encryption and safe defaults are not premium upsells.")
heading("13.2 Launch motion", 2)
for x in [
    "Recruit Tanzanian web agencies as design and distribution partners.",
    "Offer a responsible free snapshot that requires verification for continuous monitoring.",
    "Publish practical English/Kiswahili remediation guides and transparent methodology.",
    "Turn verified improvements into anonymized case studies with customer consent.",
    "Partner with hosting providers, developer communities and SME associations after product evidence exists."]:
    bullet(x)
heading("13.3 Trust centre before payment", 2)
p("Publish security overview, data handling, subprocessors, service status, vulnerability disclosure, acceptable-use policy, scan methodology, scoring methodology, limitations, contact channels and incident notification commitment.")

heading("14. Metrics, risks and governance", 1)
heading("14.1 Product scorecard", 2)
table([
    ["Dimension", "Primary measures"],
    ["Activation", "Verified domain and completed first scan within one session/day."],
    ["Value", "Findings understood, assigned and verified closed."],
    ["Retention", "Organizations with active protected assets and recurring review."],
    ["Quality", "Confirmed finding rate, disputes, failed scans and alert usefulness."],
    ["Reliability", "SLO attainment, queue delay and incident recovery."],
    ["Growth", "Design-partner conversion, agency portfolio expansion and paid retention."],
    ["Safety", "Unauthorized-scan attempts blocked, abuse incidents and security defects."],
    ["Support", "Time to first response, resolution and repeated confusion themes."]], [1.25,5.25])
heading("14.2 Principal risks", 2)
table([
    ["Risk", "Mitigation", "Early warning"],
    ["Scanner abuse", "Verification, safe profiles, limits, review and kill switch.", "Rejected private targets; unusual volume."],
    ["False confidence", "Transparent scope, evidence and limitations; no guarantee language.", "Score viewed without findings; customer disputes."],
    ["Noisy findings", "Fixtures, confidence, rule telemetry and staged rollout.", "High dispute or suppression rate."],
    ["Product compromise", "Threat model, isolation, secure SDLC and independent review.", "Secrets, auth anomalies, dependency alerts."],
    ["Weak demand", "Discovery and paid design partners before feature expansion.", "Low repeated usage or willingness to pay."],
    ["Cloud cost growth", "Quotas, caching, job budgets and cost per protected asset.", "Cost rises faster than active assets."],
    ["Regulatory error", "Local legal review, minimization and documented processing.", "Unclear retention or customer requests."],
    ["Brand collision", "Formal company/trademark search and prompt domain registration.", "Confusingly similar regional service."]], [1.35,3.25,1.90])

heading("15. Release-One completion checklist", 1)
callout("Meaning of 100%", "Release One is 100% only when every mandatory gate below is evidenced. Optional backlog size does not affect release completion.")
heading("Product and experience", 2)
for x in ["Target persona and problem validated", "English production copy complete; Kiswahili architecture ready", "All core journeys complete on mobile and desktop", "Empty, loading, error and offline-tolerant states designed", "Accessibility review passed", "Pricing and cancellation transparent"]: bullet("☐ " + x)
heading("Core capabilities", 2)
for x in ["Account, workspace and role management", "Domain ownership verification", "Safe on-demand and scheduled scans", "TLS, HTTP, DNS and email-domain checks", "Versioned findings and transparent score", "History and change detection", "Assignment, remediation and verified rescan", "Alerts and downloadable report", "Agency portfolio baseline", "Audit log and account/data controls"]: bullet("☐ " + x)
heading("AI intelligence", 2)
for x in ["Deterministic evidence remains authoritative", "Evidence-grounded summaries link to observations", "Model gateway, budgets and tenant-scoped retrieval implemented", "Prompt-injection and excessive-agency tests pass", "Human approval enforced for consequential actions", "Golden-set accuracy and bilingual evaluation accepted", "AI kill switch, rollback and provider-failure fallback proven", "Customer AI controls and retention disclosures published"]: bullet("☐ " + x)
heading("Security and privacy", 2)
for x in ["Threat model reviewed", "SSRF and redirect protections tested", "Tenant isolation tests passed", "Worker sandbox and egress controls verified", "MFA for privileged access", "Secret and dependency scanning clean", "Data inventory, retention and deletion implemented", "Vulnerability disclosure and acceptable-use policies published", "Independent security review completed", "No open Critical/High security defects"]: bullet("☐ " + x)
heading("Reliability and operations", 2)
for x in ["Production infrastructure reproducible", "Monitoring, logs, traces and alerts live", "Backups restored in rehearsal", "Incident response tabletop completed", "Status page and support process live", "SLO dashboards and runbooks approved", "Rollback and rule-disable mechanisms proven", "Cost and abuse limits configured"]: bullet("☐ " + x)
heading("Business readiness", 2)
for x in ["At least 10 active pilot organizations or an approved evidence-based threshold", "Design partners confirm usefulness", "Billing and invoicing tested", "Terms, privacy and local legal review complete", "Customer onboarding and support material complete", "Launch metrics baseline captured"]: bullet("☐ " + x)

heading("16. Post-launch continuous development", 1)
heading("16.1 Operating cadence", 2)
table([
    ["Cadence", "Activity"],
    ["Daily", "Service health, abuse signals, scan failures and security alerts."],
    ["Weekly", "Rule-quality review, customer feedback triage and release train."],
    ["Monthly", "Product metrics, cost, risk register and roadmap evidence review."],
    ["Quarterly", "Threat model update, access review, restoration exercise and strategy review."],
    ["Twice yearly", "Incident tabletop, independent security testing and policy review."],
    ["Annually", "Business continuity exercise, pricing review and regional expansion decision."]], [1.35,5.15])
heading("16.2 Candidate evolution themes", 2)
for x in [
    "Passive subdomain and certificate-transparency discovery with authorization workflow.",
    "Controlled external service exposure inventory and safe configuration checks.",
    "Brand/domain impersonation monitoring and guided takedown evidence packs.",
    "Client-side dependency and change monitoring for payment/e-commerce pages.",
    "GitHub, Jira, Linear, Slack, email and approved WhatsApp provider integrations.",
    "Partner API and webhook ecosystem.",
    "Regional languages and payment options based on measured demand.",
    "Security posture benchmarking using anonymized, consented aggregates.",
    "AI-assisted explanation and remediation only with evidence grounding, review and deterministic guardrails."]:
    bullet(x)
heading("16.3 Innovation filter", 2)
p("A proposed capability enters the roadmap only if it has a target user, verified problem, safety analysis, authorization model, measurable outcome, operational owner and sustainable cost. Novelty alone is insufficient.")

heading("17. Reference standards and glossary", 1)
heading("17.1 Standards baseline", 2)
table([
    ["Reference", "How KingaWeb uses it"],
    ["NIST CSF 2.0", "Customer risk lifecycle and organizational governance vocabulary."],
    ["OWASP ASVS", "Application-security requirements and release verification baseline."],
    ["OWASP WSTG", "Relevant web-testing methodology and evidence design."],
    ["OWASP Attack Surface Analysis", "Asset and exposure concepts."],
    ["CISA Secure by Design", "Ownership of customer outcomes, safe defaults and transparency."],
    ["OWASP Top 10 / API Top 10", "Threat awareness and engineering education—not a complete checklist."],
    ["NIST AI RMF / GenAI Profile", "Governance, mapping, measurement and management of AI risk."],
    ["OWASP GenAI Security", "Prompt injection, data disclosure, output handling and agency controls."],
    ["WCAG 2.2 AA", "Accessibility target for customer-facing experiences."],
    ["OpenTelemetry", "Vendor-neutral observability conventions."],
    ["Tanzania legal requirements", "To be mapped with qualified local counsel before public launch."]], [2.0,4.5])
heading("17.2 Glossary", 2)
for term, meaning in [
    ("Asset", "A verified public digital resource such as a domain or website."),
    ("Attack surface", "The set of reachable assets and interfaces that could be targeted."),
    ("Control plane", "The customer application, configuration and orchestration services."),
    ("Data plane", "Isolated workers that execute approved checks and collect evidence."),
    ("Finding", "A deduplicated security issue with evidence and lifecycle."),
    ("Observation", "Immutable technical evidence from one check at one time."),
    ("Rule", "Versioned logic that interprets observations."),
    ("Verified remediation", "A fix confirmed through a subsequent applicable check."),
    ("Design partner", "An early customer who provides structured, repeated product feedback."),
    ("SLO", "A measurable reliability objective for a service.")]:
    p(f"{term}. {meaning}", bold_lead=term+".")

heading("Sources consulted", 1)
sources = [
    "NIST. Cybersecurity Framework 2.0 and Small Business Quick Start resources. https://www.nist.gov/cyberframework",
    "OWASP. Web Security Testing Guide. https://wstg.owasp.org/",
    "OWASP. Attack Surface Management Top 10 and Attack Surface Analysis Cheat Sheet. https://owasp.org/",
    "CISA and international partners. Shifting the Balance of Cybersecurity Risk: Principles and Approaches for Security-by-Design and -Default.",
    "NIST. AI Risk Management Framework and Generative AI Profile. https://www.nist.gov/itl/ai-risk-management-framework",
    "OWASP. GenAI Security Project and LLM Top 10. https://genai.owasp.org/",
    "Bank of Tanzania. National Payment Systems Annual Report 2025 (context for Tanzanian digital infrastructure).",
    "Public market review of Tanzanian/African cybersecurity services and SME security-monitoring products, August 2026. Competitor capabilities must be revalidated during discovery."
]
for s in sources: bullet(s)

page_break()
heading("Approval record", 1)
table([
    ["Role", "Name", "Decision", "Date"],
    ["Product owner", "", "☐ Approve  ☐ Revise", ""],
    ["Engineering lead", "", "☐ Approve  ☐ Revise", ""],
    ["Security reviewer", "", "☐ Approve  ☐ Revise", ""],
    ["Design partner", "", "☐ Acknowledge", ""]], [1.35,1.85,2.15,1.15])

doc.core_properties.title = "KingaWeb Product & Engineering Master Plan"
doc.core_properties.subject = "Living product, security, architecture and delivery blueprint"
doc.core_properties.author = "KingaWeb Product Team"
doc.core_properties.keywords = "KingaWeb, cybersecurity, product roadmap, architecture, Tanzania, Africa"
doc.core_properties.comments = "Version 1.0 working baseline"
doc.save(OUT)
print(OUT)
